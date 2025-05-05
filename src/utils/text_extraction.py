"""
Text extraction utilities for different file formats
"""

import os
import re
import json
import csv
import xml.etree.ElementTree as ET
from io import StringIO
import logging

# Set up logging
logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX processing will be unavailable.")

try:
    import PyPDF2
    PDF_AVAILABLE = True
    PYMUPDF_AVAILABLE = False
    logger.info("PyPDF2 is available for PDF processing.")
except ImportError:
    PDF_AVAILABLE = False
    PYMUPDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed.")

# Try PyMuPDF separately
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
    PYMUPDF_AVAILABLE = True
    logger.info("PyMuPDF is available for PDF processing.")
except ImportError:
    if not PDF_AVAILABLE:
        PDF_AVAILABLE = False
        PYMUPDF_AVAILABLE = False
        logger.warning("Neither PyPDF2 nor PyMuPDF installed. PDF processing will be unavailable.")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.warning("pandas not installed. Excel processing will be unavailable.")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not installed. Excel processing will be unavailable.")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    logger.warning("beautifulsoup4 not installed. HTML processing will be unavailable.")

def extract_text_from_file(file_path):
    """
    Extract text from a file based on its extension
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    try:
        if ext == '.txt':
            return extract_text_from_txt(file_path)
        elif ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return extract_text_from_docx(file_path)
        elif ext == '.doc':
            logger.warning(f"Legacy .doc format not directly supported: {file_path}")
            return f"[Legacy .doc format not directly supported: {file_path}]"
        elif ext in ['.csv', '.tsv']:
            return extract_text_from_csv(file_path)
        elif ext in ['.json']:
            return extract_text_from_json(file_path)
        elif ext in ['.xml']:
            return extract_text_from_xml(file_path)
        elif ext in ['.html', '.htm']:
            return extract_text_from_html(file_path)
        elif ext in ['.xlsx', '.xls']:
            return extract_text_from_excel(file_path)
        elif ext in ['.md', '.markdown']:
            return extract_text_from_markdown(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return f"[Unsupported file format: {ext}]"
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return f"[Error extracting text: {str(e)}]"

def extract_text_from_txt(file_path):
    """
    Extract text from a plain text file
    
    Args:
        file_path (str): Path to the text file
        
    Returns:
        str: Extracted text
    """
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()

def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text
    """
    if not PDF_AVAILABLE:
        return "[PDF processing not available. Install PyPDF2 or PyMuPDF.]"
    
    text = ""
    
    # Try PyMuPDF first if available (better quality)
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            logger.error(f"PyMuPDF error for {file_path}: {str(e)}")
            # Fall back to PyPDF2
    
    # Use PyPDF2 as fallback
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PyPDF2 error for {file_path}: {str(e)}")
        return f"[Error extracting PDF text: {str(e)}]"

def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text
    """
    if not DOCX_AVAILABLE:
        return "[DOCX processing not available. Install python-docx.]"
    
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"DOCX error for {file_path}: {str(e)}")
        return f"[Error extracting DOCX text: {str(e)}]"

def extract_text_from_csv(file_path):
    """
    Extract text from a CSV file
    
    Args:
        file_path (str): Path to the CSV file
        
    Returns:
        str: Extracted text
    """
    try:
        # Detect delimiter
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            sample = f.read(1024)
            dialect = csv.Sniffer().sniff(sample)
            delimiter = dialect.delimiter
        
        # Read CSV
        rows = []
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f, delimiter=delimiter)
            headers = next(reader, [])
            rows.append(" | ".join(headers))
            
            for row in reader:
                rows.append(" | ".join(row))
        
        return "\n".join(rows)
    except Exception as e:
        logger.error(f"CSV error for {file_path}: {str(e)}")
        return f"[Error extracting CSV text: {str(e)}]"

def extract_text_from_json(file_path):
    """
    Extract text from a JSON file
    
    Args:
        file_path (str): Path to the JSON file
        
    Returns:
        str: Extracted text
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            data = json.load(f)
        
        # Convert JSON to string representation
        return json.dumps(data, indent=2)
    except Exception as e:
        logger.error(f"JSON error for {file_path}: {str(e)}")
        return f"[Error extracting JSON text: {str(e)}]"

def extract_text_from_xml(file_path):
    """
    Extract text from an XML file
    
    Args:
        file_path (str): Path to the XML file
        
    Returns:
        str: Extracted text
    """
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # Extract all text content
        text_parts = []
        
        def extract_text_from_element(element):
            if element.text and element.text.strip():
                text_parts.append(f"{element.tag}: {element.text.strip()}")
            
            for child in element:
                extract_text_from_element(child)
        
        extract_text_from_element(root)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"XML error for {file_path}: {str(e)}")
        return f"[Error extracting XML text: {str(e)}]"

def extract_text_from_html(file_path):
    """
    Extract text from an HTML file
    
    Args:
        file_path (str): Path to the HTML file
        
    Returns:
        str: Extracted text
    """
    if not BS4_AVAILABLE:
        return "[HTML processing not available. Install beautifulsoup4.]"
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text()
        
        # Break into lines and remove leading and trailing space on each
        lines = (line.strip() for line in text.splitlines())
        
        # Break multi-headlines into a line each
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        
        # Drop blank lines
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        logger.error(f"HTML error for {file_path}: {str(e)}")
        return f"[Error extracting HTML text: {str(e)}]"

def extract_text_from_excel(file_path):
    """
    Extract text from an Excel file
    
    Args:
        file_path (str): Path to the Excel file
        
    Returns:
        str: Extracted text
    """
    if not PANDAS_AVAILABLE or not OPENPYXL_AVAILABLE:
        return "[Excel processing not available. Install pandas and openpyxl.]"
    
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names
        
        text_parts = []
        
        for sheet_name in sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Add sheet name
            text_parts.append(f"Sheet: {sheet_name}")
            
            # Convert to CSV-like format
            csv_buffer = StringIO()
            df.to_csv(csv_buffer, index=False)
            csv_text = csv_buffer.getvalue()
            
            text_parts.append(csv_text)
            text_parts.append("")  # Add separator between sheets
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Excel error for {file_path}: {str(e)}")
        return f"[Error extracting Excel text: {str(e)}]"

def extract_text_from_markdown(file_path):
    """
    Extract text from a Markdown file
    
    Args:
        file_path (str): Path to the Markdown file
        
    Returns:
        str: Extracted text
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        
        # Remove Markdown formatting (basic)
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Remove emphasis
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        
        # Remove inline code
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        # Remove links
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        
        # Remove images
        text = re.sub(r'!\[(.*?)\]\(.*?\)', r'\1', text)
        
        return text
    except Exception as e:
        logger.error(f"Markdown error for {file_path}: {str(e)}")
        return f"[Error extracting Markdown text: {str(e)}]"

def clean_extracted_text(text):
    """
    Clean extracted text
    
    Args:
        text (str): Text to clean
        
    Returns:
        str: Cleaned text
    """
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove extra newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Normalize spaces
    text = re.sub(r' {2,}', ' ', text)
    
    return text.strip()
